import os
import PyPDF2
import docx
from typing import List, Dict, Any, Optional
from app.models.document import Document, DocumentMetadata


class DocumentLoader:
    """다양한 형식의 문서를 로드하는 유틸리티 클래스"""

    @staticmethod
    def load_pdf(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """PDF 파일을 로드하여 Document 객체로 변환"""
        text = ""

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"

        if metadata is None:
            metadata = {}

        # 기본 메타데이터 추가
        doc_metadata = {
            "source": file_path,
            "title": os.path.basename(file_path),
            "document_type": "pdf",
            **metadata
        }

        return Document(content=text, metadata=doc_metadata)

    @staticmethod
    def load_docx(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Word 문서를 로드하여 Document 객체로 변환"""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        if metadata is None:
            metadata = {}

        # 기본 메타데이터 추가
        doc_metadata = {
            "source": file_path,
            "title": os.path.basename(file_path),
            "document_type": "docx",
            **metadata
        }

        return Document(content=text, metadata=doc_metadata)

    @staticmethod
    def load_text(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """텍스트 파일을 로드하여 Document 객체로 변환"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        if metadata is None:
            metadata = {}

        # 기본 메타데이터 추가
        doc_metadata = {
            "source": file_path,
            "title": os.path.basename(file_path),
            "document_type": "text",
            **metadata
        }

        return Document(content=text, metadata=doc_metadata)

    @classmethod
    def load_document(cls, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """파일 확장자에 따라 적절한 로더를 선택하여 문서 로드"""
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            return cls.load_pdf(file_path, metadata)
        elif file_extension in ['.docx', '.doc']:
            return cls.load_docx(file_path, metadata)
        elif file_extension in ['.txt', '.md']:
            return cls.load_text(file_path, metadata)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_extension}")
